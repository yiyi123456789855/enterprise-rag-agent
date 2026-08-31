from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from app.types import Paragraph


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass(slots=True)
class ParsedDocument:
    paragraphs: list[Paragraph]
    title: str | None = None


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")
    if suffix in {".txt", ".md"}:
        return _parse_text(content, markdown=suffix == ".md")
    if suffix == ".pdf":
        return _parse_pdf(content)
    return _parse_docx(content)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("The text file encoding is not UTF-8 or GB18030")


def _parse_text(content: bytes, *, markdown: bool) -> ParsedDocument:
    text = _decode_text(content).replace("\r\n", "\n")
    current_heading: str | None = None
    buffer: list[str] = []
    paragraphs: list[Paragraph] = []

    def flush() -> None:
        if buffer:
            joined = "\n".join(buffer).strip()
            if joined:
                paragraphs.append(Paragraph(text=joined, heading=current_heading))
            buffer.clear()

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if markdown and re.match(r"^#{1,6}\s+", line):
            flush()
            current_heading = re.sub(r"^#{1,6}\s+", "", line).strip()
        elif not line:
            flush()
        else:
            buffer.append(line)
    flush()
    title = next((item.heading for item in paragraphs if item.heading), None)
    return ParsedDocument(paragraphs=paragraphs, title=title)


def _parse_pdf(content: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency error is user-facing
        raise RuntimeError("PDF parsing requires pypdf") from exc

    reader = PdfReader(BytesIO(content))
    paragraphs: list[Paragraph] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        for block in re.split(r"\n\s*\n|(?<=。)\s*\n", page_text):
            normalized = re.sub(r"[ \t]+", " ", block).strip()
            if normalized:
                paragraphs.append(Paragraph(text=normalized, page_number=page_number))
    return ParsedDocument(paragraphs=paragraphs)


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency error is user-facing
        raise RuntimeError("DOCX parsing requires python-docx") from exc

    document = Document(BytesIO(content))
    paragraphs: list[Paragraph] = []
    current_heading: str | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading") or style_name.startswith("标题"):
            current_heading = text
            continue
        paragraphs.append(Paragraph(text=text, heading=current_heading))
    return ParsedDocument(paragraphs=paragraphs, title=current_heading)

